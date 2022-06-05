from typing import List
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def load_text_from_file(filename: str="text_to_process.txt") -> List[str]:
    with open(filename, "r", encoding="utf-8") as input_file:
        data = input_file.readlines()
    if data[-1][-1] == "\n":
        data[-1] = data[-1][:-1]
    return data

def process_data_to_txt(data: List[str], additional_spaces_between_paragraphs: bool=False, tabulations_before_each_paragraph: bool=False):
    processed_text = ""
    for line_number, line_content in enumerate(data):
        if line_number == 0:
            previous_line = line_content
            previous_broken_line = False
            continue
        current_line = line_content
        if (broken_line := current_line[0].isalpha() and current_line[0].islower()):
            previous_line = previous_line[:-1] + " "
        additional_paragraph_for_txt = "\n" if not broken_line and additional_spaces_between_paragraphs else ""
        tab_character = "\t" if not previous_broken_line and tabulations_before_each_paragraph else ""
        processed_text += tab_character + previous_line + additional_paragraph_for_txt
        previous_line = current_line
        previous_broken_line = broken_line
    processed_text += previous_line
    return processed_text[:-1] if processed_text[-1] == "\n" else processed_text

def process_data_to_word(data: List[str], additional_spaces_between_paragraphs: bool=False, tabulations_before_each_paragraph: bool=False, justify_paragraph: bool=True):
    document = Document()
    processed_paragraph = ""
    for line_number in range((number_of_lines := len(data)) - 1):
        if line_number == 0:
            paragraph = document.add_paragraph()
        if data[line_number] == "\n":
            continue
        if data[line_number + 1][0].isalpha() and data[line_number + 1][0].islower():
            processed_paragraph += data[line_number][:-1] + " "
            if line_number == number_of_lines - 2:
                processed_paragraph += data[-1][:-1]
        else:
            processed_paragraph += data[line_number][:-1]
            processed_paragraph = "\t" + processed_paragraph if tabulations_before_each_paragraph else processed_paragraph
            if justify_paragraph:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.add_run(processed_paragraph)
            if additional_spaces_between_paragraphs and line_number != len(data) - 1:
                document.add_paragraph()
            paragraph = document.add_paragraph()
            processed_paragraph = ""
    processed_paragraph += data[-1]
    processed_paragraph = "\t" + processed_paragraph if tabulations_before_each_paragraph else processed_paragraph
    if justify_paragraph:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.add_run(processed_paragraph)
    return document

def save_data_to_txt(processed_text: str):
    with open("processed_text.txt", "w", encoding="utf-8") as out_file_txt:
        out_file_txt.write(processed_text)

def save_data_to_word(document: Document):
    document.save("processed_text.docx")

if __name__ == '__main__':
    text_to_process = load_text_from_file()
    processed_text_txt = process_data_to_txt(text_to_process)
    processed_document_word = process_data_to_word(text_to_process)
    save_data_to_txt(processed_text_txt)
    save_data_to_word(processed_document_word)
