from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

number_of_lines_in_each_paragraph = [1, 1, 1, 7, 6, 10, 8, 9]

make_additional_spaces_between_paragraphs = False
make_tabulations_at_the_beginning_of_the_paragraph = False
fully_justify_paragraph = True

target_filename = "text_to_process.txt"
document = Document()


def save_as_new_file(processed_text):
    with open('processed_text.txt', 'w', encoding='utf-8') as out_file_txt:
        out_file_txt.write(processed_text)
    document.save('processed_text.docx')


def copy_and_fix_text_from_file(input_filename):
    processed_text_for_txt = ""
    count_lines = 1
    count_paragraphs = 0
    try:
        with open(input_filename, 'r', encoding='utf-8') as in_file:
            for (line_number, line_content) in enumerate(in_file):
                temporary_line = line_content
                if count_lines == 1:
                    if temporary_line[0].isalpha():
                        if temporary_line[0].islower():
                            raise SyntaxError(f"Possible mistake in counting the lines. "
                                              f"Line, that mistake occurred: {line_number + 1}")
                    paragraph_for_word = document.add_paragraph()
                    processed_text_for_word = ""
                    if make_tabulations_at_the_beginning_of_the_paragraph:
                        paragraph_for_word.paragraph_format.left_indent = Inches(0.5)
                        processed_text_for_txt += '\t'
                if count_lines != number_of_lines_in_each_paragraph[count_paragraphs]:
                    temporary_line = temporary_line.replace("\n", " ")
                    processed_text_for_txt += temporary_line
                    processed_text_for_word += temporary_line
                    count_lines += 1
                else:
                    processed_text_for_txt += temporary_line
                    temporary_line = temporary_line.replace("\n", "")
                    processed_text_for_word += temporary_line
                    if fully_justify_paragraph:
                        paragraph_for_word.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    paragraph_for_word.add_run(processed_text_for_word)
                    if make_additional_spaces_between_paragraphs:
                        processed_text_for_txt += '\n'
                        document.add_paragraph()
                    count_lines = 1
                    count_paragraphs += 1
    except OSError:
        raise FileNotFoundError(f'File "{input_filename}" could not be opened. Check file name or file path.')
    processed_text_for_txt = processed_text_for_txt[:-1]
    return processed_text_for_txt


if __name__ == '__main__':
    processed_text = copy_and_fix_text_from_file(target_filename)
    save_as_new_file(processed_text)
