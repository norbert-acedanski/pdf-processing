import sys
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

numberOfLinesInEachParagraph = [1, 1, 1, 7, 6, 10, 8, 9]

makeAdditionalSpacesBetweenParagraphs = True
makeTabulationsAtTheBeginningOfTheParagraph = False
fullyJustifyParagraph = True

targetFileName = "textToProcess.txt"
document = Document()

def saveAsNewFile(processedText):
    with open('processedText.txt', 'w', encoding='utf-8') as outFile_TXT:
        outFile_TXT.write(processedText)
    document.save('processedText.docx')

def copyAndFixTextFromFile(inputFileName):
    processedTextForTxt = ""
    countLines = 1
    countParagraphs = 0
    try:
        with open(inputFileName, 'r', encoding='utf-8') as inFile:
            for (lineNumber, lineContent) in enumerate(inFile):
                temporaryLine = lineContent
                if countLines == 1:
                    if temporaryLine[0].isalpha():
                        if temporaryLine[0].islower():
                            print("Possible mistake in counting the lines. Line, that mistake occurred: " + str(lineNumber + 1))
                            sys.exit()
                    paragraphForWord = document.add_paragraph()
                    processedTextForWord = ""
                    if makeTabulationsAtTheBeginningOfTheParagraph:
                        paragraphForWord.paragraph_format.left_indent = Inches(0.5)
                        processedTextForTxt += '\t'
                if countLines != numberOfLinesInEachParagraph[countParagraphs]:
                    temporaryLine = temporaryLine.replace("\n", " ")
                    processedTextForTxt += temporaryLine
                    processedTextForWord += temporaryLine
                    countLines += 1
                else:
                    processedTextForTxt += temporaryLine
                    temporaryLine = temporaryLine.replace("\n", "")
                    processedTextForWord += temporaryLine
                    if fullyJustifyParagraph:
                        paragraphForWord.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    paragraphForWord.add_run(processedTextForWord)
                    if makeAdditionalSpacesBetweenParagraphs:
                        processedTextForTxt += '\n'
                        document.add_paragraph()
                    countLines = 1
                    countParagraphs += 1
    except OSError:
        print("File \"" + inputFileName + "\" could not be opened. Check file name or file path.")
        sys.exit()
    processedTextForTxt = processedTextForTxt[:-1]
    return processedTextForTxt

if __name__ == '__main__':
    processedText = copyAndFixTextFromFile(targetFileName)
    saveAsNewFile(processedText)
